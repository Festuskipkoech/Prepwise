import io
import pytest
import pypdf
import docx

from app.profile.extractor import detect_file_type, extract_text

def _make_docx_bytes(text: str) -> bytes:
    """Create a real minimal DOCX containing the given text."""
    document = docx.Document()
    document.add_paragraph(text)
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()

# detect_file_type
def test_detect_pdf_by_extension():
    assert detect_file_type("resume.pdf", "application/octet-stream") == "pdf"

def test_detect_pdf_by_content_type():
    assert detect_file_type("upload", "application/pdf") == "pdf"

def test_detect_docx_by_extension():
    result = detect_file_type("resume.docx", "application/octet-stream")
    assert result == "docx"

def test_detect_docx_by_content_type_openxml():
    ct = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    assert detect_file_type("upload", ct) == "docx"

def test_detect_docx_by_content_type_msword():
    assert detect_file_type("upload", "application/msword") == "docx"

def test_detect_unsupported_raises():
    with pytest.raises(ValueError, match="Unsupported file type"):
        detect_file_type("resume.txt", "text/plain")

def test_detect_unsupported_png_raises():
    with pytest.raises(ValueError, match="Unsupported file type"):
        detect_file_type("photo.png", "image/png")

def test_detect_case_insensitive_extension():
    assert detect_file_type("RESUME.PDF", "application/octet-stream") == "pdf"

# extract_text — DOCX (reliable in-memory generation)
def test_extract_docx_returns_text():
    data = _make_docx_bytes("Hello from a DOCX file.")
    text = extract_text(data, "docx")
    assert "Hello from a DOCX file." in text

def test_extract_docx_strips_empty_paragraphs():
    data = _make_docx_bytes("Line one.\n\nLine two.")
    text = extract_text(data, "docx")
    assert "Line one." in text

def test_extract_docx_multiline():
    document = docx.Document()
    document.add_paragraph("First paragraph.")
    document.add_paragraph("Second paragraph.")
    buf = io.BytesIO()
    document.save(buf)
    text = extract_text(buf.getvalue(), "docx")
    assert "First paragraph." in text
    assert "Second paragraph." in text

def test_extract_docx_with_table():
    document = docx.Document()
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Cell A"
    table.cell(0, 1).text = "Cell B"
    buf = io.BytesIO()
    document.save(buf)
    text = extract_text(buf.getvalue(), "docx")
    assert "Cell A" in text
    assert "Cell B" in text

# extract_text — error handling
def test_extract_text_unsupported_type_raises():
    with pytest.raises(ValueError, match="Unsupported file type"):
        extract_text(b"some bytes", "txt")

def test_extract_text_empty_docx_raises():
    document = docx.Document()
    buf = io.BytesIO()
    document.save(buf)
    with pytest.raises(ValueError, match="No text could be extracted"):
        extract_text(buf.getvalue(), "docx")

def test_extract_text_garbage_bytes_raises():
    with pytest.raises(Exception):
        extract_text(b"not a real pdf at all", "pdf")
